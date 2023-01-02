import os

import PyPDF2
from django.shortcuts import render, redirect
from django import forms
from docx import Document
from cv_app import models
import hashlib


def md5(password):
    md5_obj = hashlib.md5()
    md5_obj.update(password.encode('utf-8'))
    return md5_obj.hexdigest()


# Create your views here.
def signup(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")
    if request.method == "POST":
        name = request.POST.get("name")
        email = request.POST.get("email")
        password = md5(request.POST.get("password"))
        print(name, email, password)
        models.UserInfo.objects.create(name=name, email=email, pwd=password)
        return redirect("../success")
    return render(request, "signup.html")


class LoginForm(forms.Form):
    email = forms.CharField(label="email", widget=forms.EmailInput(
        attrs={"class": "form-control", "placeholder": "Email address", "required": "required"}))
    pwd = forms.CharField(label="password", widget=forms.PasswordInput(
        attrs={"class": "form-control", "placeholder": "Password", "required": "required"}))


def login(request):
    if request.method == "GET":
        form = LoginForm()
        return render(request, 'login.html', {'form': form})
    form = LoginForm(data=request.POST)
    if form.is_valid():
        admin = models.UserInfo.objects.filter(email=form.cleaned_data['email'],
                                               pwd=md5(form.cleaned_data['pwd'])).first()
        if not admin:
            return redirect("../error_pwd")
        request.session["info"] = admin.name
        return redirect('../index')
    return render(request, 'login.html', {'form': form})


def index(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")
    query = models.CV.objects.filter(if_follow=1)
    return render(request, "index.html", {'query': query})


def index_delete(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")
    nid = request.GET.get('nid')
    models.CV.objects.filter(id=nid).update(if_follow=0)
    return redirect('/index')


class FileFieldForm(forms.Form):
    file_field = forms.FileField(widget=forms.ClearableFileInput(attrs={'multiple': True, 'class': "bg-info"}))


def single_upload(f):
    file_path = os.path.join(os.path.abspath('../cv_app/static/'), 'uploaded_cv', f.name)
    with open(file_path, 'wb+') as destination:  # 写文件word
        for chunk in f.chunks():
            destination.write(chunk)
    destination.close()
    return file_path


def scratch(path):
    content = ''
    if path.endswith(('.doc', '.docx')):
        document = Document(path)
        tables = document.tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        content += paragraph.text
        for para in document.paragraphs:
            content += para.text
    elif path.endswith('.pdf'):
        pdffile = open(path, 'rb')  # 读取pdf文件
        pdfreader = PyPDF2.PdfFileReader(pdffile)
        for i in range(0, pdfreader.numPages):
            content += pdfreader.getPage(i).extractText()
    content = content.replace('\n', '').replace(' ', '')
    return content


def upload(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")
    if request.method == 'POST':
        form = FileFieldForm(request.POST, request.FILES)
        files = request.FILES.getlist('file_field')  # 获得多个文件上传进来的文件列表。
        if form.is_valid():  # 表单数据如果合法
            for f in files:
                file_path = single_upload(f)  # 处理上传来的文件
                content = scratch(file_path)
                models.CV.objects.create(name=f.name, content=content)
            return redirect('../manage')
        return redirect('/error')
    else:
        form = FileFieldForm()
    return render(request, "upload.html", locals())


def manage(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")
    if request.method == "POST":
        keywords = request.POST.get("keywords").split('+')
        all_query_sql = "select * from cv_app_cv where (content like '%" + keywords[0] + "%' or name like '%" + \
                        keywords[0] + "%')"
        if len(keywords) > 1:
            for i in keywords[1:]:
                all_query_sql += " and (content like '%" + i + "%' or name like '%" + i + "%')"
        all_query = models.CV.objects.raw(all_query_sql)
        unfollowed_query_sql = all_query_sql + ' and (if_follow=0)'
        unfollowed_query = models.CV.objects.raw(unfollowed_query_sql)
        followed_query_sql = all_query_sql + ' and (if_follow=1)'
        followed_query = models.CV.objects.raw(followed_query_sql)
        return render(request, "manage.html",
                      {'all_query': all_query, 'unfollowed_query': unfollowed_query, 'followed_query': followed_query})

    all_query = models.CV.objects.all()
    unfollowed_query = models.CV.objects.filter(if_follow=0)
    followed_query = models.CV.objects.filter(if_follow=1)
    return render(request, "manage.html",
                  {'all_query': all_query, 'unfollowed_query': unfollowed_query, 'followed_query': followed_query})


def manage_delete(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")
    nid = request.GET.get('nid')
    models.CV.objects.filter(id=nid).update(if_follow=0)
    return redirect('/manage')


def manage_follow(request):
    info = request.session.get("info")
    if not info:
        return redirect("../login")

    nid = request.GET.get('nid')
    models.CV.objects.filter(id=nid).update(if_follow=1)
    return redirect('/manage')


def error(request):
    return render(request, "error.html")


def success(request):
    return render(request, "success.html")


def logout(request):
    request.session.clear()
    return redirect("../login")


def error_pwd(request):
    return render(request, "error_pwd.html")
